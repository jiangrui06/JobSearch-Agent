"""简历解析模块：支持 PDF、DOCX、TXT 格式简历的文本提取。"""

import logging
from pathlib import Path
from typing import Union

logger = logging.getLogger(__name__)


def parse_resume(file_path: Union[str, Path]) -> str:
    """解析简历文件，提取文本内容。

    Args:
        file_path: 简历文件的路径

    Returns:
        简历全文文本

    Raises:
        FileNotFoundError: 文件不存在时
        ValueError: 不支持的格式或解析失败
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"简历文件未找到: {file_path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(path)
    elif suffix == ".docx":
        return _parse_docx(path)
    elif suffix in (".doc", ".dot"):
        return _parse_doc_fallback(path)
    elif suffix == ".txt":
        return path.read_text(encoding="utf-8")
    else:
        raise ValueError(f"不支持的简历格式: {suffix}（支持: .pdf, .docx, .txt）")


def _parse_pdf(path: Path) -> str:
    """使用 pdfplumber 提取 PDF 文本。"""
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber 未安装，尝试回退到 PyPDF2...")
        return _parse_pdf_fallback(path)

    text_parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    full_text = "\n".join(text_parts).strip()
    if not full_text:
        raise ValueError("PDF 解析结果为空，请检查文件是否可读或尝试文本格式简历。")
    logger.info(f"成功解析简历 PDF（共 {len(text_parts)} 页，{len(full_text)} 字符）")
    return full_text


def _parse_pdf_fallback(path: Path) -> str:
    """使用 PyPDF2 作为 pdfplumber 的回退方案。"""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ImportError(
            "请安装 PDF 解析依赖：pip install pdfplumber 或 PyPDF2"
        )

    reader = PdfReader(path)
    text_parts = [
        page.extract_text() for page in reader.pages if page.extract_text()
    ]
    full_text = "\n".join(text_parts).strip()
    if not full_text:
        raise ValueError("PDF 解析结果为空（PyPDF2），请检查文件。")
    logger.info(f"成功解析简历 PDF（PyPDF2，共 {len(text_parts)} 页）")
    return full_text


def _parse_docx(path: Path) -> str:
    """使用 python-docx 提取 .docx 文件文本。"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("需要安装 python-docx：pip install python-docx")

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # 也尝试提取表格内容
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_texts.append(" | ".join(cells))

    full_text = "\n".join(paragraphs + table_texts).strip()
    if not full_text:
        raise ValueError("DOCX 解析结果为空，请检查文件。")
    logger.info(f"成功解析简历 DOCX（{len(paragraphs)} 段落，{len(full_text)} 字符）")
    return full_text


def _parse_doc_fallback(path: Path) -> str:
    """尝试用 python-docx 打开 .doc 文件（部分 .doc 可读）。

    较旧的 .doc 格式无法直接用 python-docx 解析，给出提示。
    """
    # 先尝试用 python-docx 打开（部分 .doc 文件实际是 docx 格式）
    try:
        return _parse_docx(path)
    except Exception:
        pass

    # 尝试用 antiword 或 catdoc 命令行工具（如果安装）
    import subprocess
    for cmd in (["antiword", str(path)], ["catdoc", str(path)]):
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
            if result.returncode == 0 and result.stdout.strip():
                logger.info(f"通过 {cmd[0]} 解析 .doc 成功")
                return result.stdout.strip()
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    raise ValueError(
        "无法解析 .doc 格式。建议将文件另存为 .docx 或 .pdf 格式后重试。\n"
        "如需直接支持 .doc，请安装 antiword 或 catdoc。"
    )
